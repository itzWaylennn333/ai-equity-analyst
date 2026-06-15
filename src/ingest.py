"""Document ingestion (Layer 1 of the platform).

Parse user-uploaded documents (PDF / HTML / DOCX / XLSX / CSV / TXT) into clean,
classified, provenance-tagged, structure-aware chunks that downstream layers
(RAG, extraction agents, sentiment) consume. No network; pure local parsing.

Public API:
    ingest_file(cfg, ticker, path)  -> {doc_type, n_chunks, ...}
    ingest_files(cfg, ticker, paths) -> list of the above; writes chunks.jsonl + provenance
    load_chunks(cfg, ticker)        -> list[dict]   (for RAG / display)
"""
from __future__ import annotations

import json
import re
import hashlib
from pathlib import Path

import pandas as pd

import utils

# --------------------------------------------------------------------------- #
# Parsers (each returns {"text": str, "tables": list[DataFrame]})
# --------------------------------------------------------------------------- #
def _parse_pdf(path: Path) -> dict:
    text_parts, tables = [], []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
                for t in (page.extract_tables() or []):
                    try:
                        tables.append(pd.DataFrame(t[1:], columns=t[0]))
                    except Exception:
                        pass
    except Exception:
        # fallback to PyMuPDF text-only
        import fitz
        doc = fitz.open(path)
        text_parts = [pg.get_text() for pg in doc]
    return {"text": "\n".join(text_parts), "tables": tables}


def _parse_html(path: Path) -> dict:
    from bs4 import BeautifulSoup
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    text = re.sub(r"\n{3,}", "\n\n", soup.get_text("\n"))
    tables = []
    try:
        tables = pd.read_html(path)  # best-effort
    except Exception:
        pass
    return {"text": text, "tables": tables}


def _parse_docx(path: Path) -> dict:
    import docx
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs)
    tables = []
    for tbl in d.tables:
        rows = [[c.text for c in r.cells] for r in tbl.rows]
        if rows:
            tables.append(pd.DataFrame(rows[1:], columns=rows[0]))
    return {"text": text, "tables": tables}


def _parse_xlsx(path: Path) -> dict:
    sheets = pd.read_excel(path, sheet_name=None)
    parts, tables = [], []
    for name, df in sheets.items():
        parts.append(f"# Sheet: {name}\n{df.to_string(max_rows=200)}")
        tables.append(df)
    return {"text": "\n\n".join(parts), "tables": tables}


def _parse_csv(path: Path) -> dict:
    df = pd.read_csv(path)
    return {"text": f"# {path.name}\n{df.to_string(max_rows=500)}", "tables": [df]}


def _parse_txt(path: Path) -> dict:
    return {"text": path.read_text(encoding="utf-8", errors="ignore"), "tables": []}


_PARSERS = {".pdf": _parse_pdf, ".htm": _parse_html, ".html": _parse_html,
            ".docx": _parse_docx, ".xlsx": _parse_xlsx, ".xls": _parse_xlsx,
            ".csv": _parse_csv, ".txt": _parse_txt}


# --------------------------------------------------------------------------- #
# Classification
# --------------------------------------------------------------------------- #
def classify(text: str, filename: str) -> str:
    # Scan a generous slice (EDGAR inline-XBRL HTML buries readable text after metadata).
    t, fn = text[:60000].lower(), filename.lower()
    if (re.search(r"10[-_ ]?k|10[-_ ]?q|8[-_ ]?k", fn)
            or re.search(r"\bform\s*10-?[kq]\b|annual report on form|\b10-k\b|\b10-q\b", t)):
        return "filing"
    if re.search(r"earnings call|conference call|prepared remarks|transcript|q&a session", t + " " + fn):
        return "transcript"
    if fn.endswith((".xlsx", ".xls")) or "valuation model" in t:
        return "model"
    if re.search(r"investor presentation|investor day|earnings presentation", t + " " + fn):
        return "presentation"
    if re.search(r"press release|for immediate release|reuters|bloomberg|analyst note", t + " " + fn):
        return "news"
    return "other"


# --------------------------------------------------------------------------- #
# Chunking (structure-aware for filings)
# --------------------------------------------------------------------------- #
_ITEM_RE = re.compile(r"(?im)^\s*(item\s+\d+[a-z]?\.?.*)$")


def _split_sections(text: str, doc_type: str) -> list[tuple[str, str]]:
    """Return [(section_label, section_text)]. Filings split on SEC 'Item N.' headers."""
    if doc_type != "filing":
        return [("body", text)]
    matches = list(_ITEM_RE.finditer(text))
    if not matches:
        return [("body", text)]
    sections = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        label = re.sub(r"\s+", " ", m.group(1)).strip()[:80]
        sections.append((label, text[start:end]))
    if matches[0].start() > 0:
        sections.insert(0, ("preamble", text[: matches[0].start()]))
    return sections


def _window(words: list[str], max_words: int, overlap: float) -> list[tuple[int, list[str]]]:
    step = max(1, int(max_words * (1 - overlap)))
    out = []
    for i in range(0, len(words), step):
        out.append((i, words[i:i + max_words]))
        if i + max_words >= len(words):
            break
    return out


def chunk_document(parsed: dict, doc_type: str, ticker: str, source_file: str,
                   max_words: int = 600, overlap: float = 0.15) -> list[dict]:
    chunks = []
    for section, sec_text in _split_sections(parsed["text"], doc_type):
        words = sec_text.split()
        if not words:
            continue
        for wi, win in _window(words, max_words, overlap):
            body = " ".join(win)
            cid = hashlib.sha1(f"{ticker}|{source_file}|{section}|{wi}".encode()).hexdigest()[:16]
            chunks.append({
                "chunk_id": cid, "ticker": ticker, "source_file": source_file,
                "doc_type": doc_type, "section": section, "word_start": wi,
                "n_words": len(win), "text": body,
            })
    return chunks


# --------------------------------------------------------------------------- #
# Orchestration
# --------------------------------------------------------------------------- #
def _doc_dir(cfg: dict, ticker: str) -> Path:
    d = utils.resolve(cfg["data"]["processed_dir"]) / ticker / "documents"
    d.mkdir(parents=True, exist_ok=True)
    return d


def ingest_file(cfg: dict, ticker: str, path: str | Path) -> dict:
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in _PARSERS:
        return {"source_file": path.name, "error": f"unsupported type {ext}", "chunks": []}
    parsed = _PARSERS[ext](path)
    doc_type = classify(parsed["text"], path.name)
    chunks = chunk_document(parsed, doc_type, ticker, path.name)
    return {"source_file": path.name, "doc_type": doc_type, "n_tables": len(parsed["tables"]),
            "n_chunks": len(chunks), "n_chars": len(parsed["text"]), "chunks": chunks}


def ingest_files(cfg: dict, ticker: str, paths: list[str | Path]) -> list[dict]:
    ddir = _doc_dir(cfg, ticker)
    results, all_chunks = [], []
    for p in paths:
        r = ingest_file(cfg, ticker, p)
        all_chunks.extend(r.pop("chunks"))
        results.append(r)
    # write/append the chunk store (JSONL) for this ticker
    out = ddir / "chunks.jsonl"
    with open(out, "w", encoding="utf-8") as f:
        for c in all_chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")
    utils.record_provenance(cfg["data"]["manifest"], "upload", ticker,
                            url="local-upload", files=[str(out)],
                            extra={"docs": [r["source_file"] for r in results], "n_chunks": len(all_chunks)})
    return results


def load_chunks(cfg: dict, ticker: str) -> list[dict]:
    out = _doc_dir(cfg, ticker) / "chunks.jsonl"
    if not out.exists():
        return []
    return [json.loads(line) for line in out.read_text(encoding="utf-8").splitlines() if line.strip()]


if __name__ == "__main__":
    import sys
    sys.path.insert(0, str(utils.PROJECT_ROOT / "src"))
    import registry
    tk = sys.argv[1] if len(sys.argv) > 1 else "PYPL"
    files = sys.argv[2:]
    cfg = registry.load_company(tk)
    for r in ingest_files(cfg, tk, files):
        print(r)
